"""
Service for parsing CV files (PDF and DOCX).
"""

import os
import logging
from typing import Tuple
import fitz  # PyMuPDF
from docx import Document

logger = logging.getLogger(__name__)


class CVParser:
    """Parser for CV files in PDF and DOCX formats."""
    
    SUPPORTED_FORMATS = {'.pdf', '.docx'}
    
    @staticmethod
    def is_supported_format(file_path: str) -> bool:
        """
        Check if file format is supported.
        
        Args:
            file_path: Path to file
            
        Returns:
            True if format is supported
        """
        _, ext = os.path.splitext(file_path)
        return ext.lower() in CVParser.SUPPORTED_FORMATS
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """
        Extract text from PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        try:
            text = ""
            with fitz.open(file_path) as pdf:
                for page_num in range(len(pdf)):
                    page = pdf[page_num]
                    text += page.get_text()
            
            if not text.strip():
                logger.warning(f"No text extracted from PDF: {file_path}")
            
            return text
        
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text
        """
        try:
            text = ""
            doc = Document(file_path)
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
            
            if not text.strip():
                logger.warning(f"No text extracted from DOCX: {file_path}")
            
            return text
        
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    @staticmethod
    def extract_text(file_path: str) -> str:
        """
        Extract text from CV file (automatically detects format).
        
        Args:
            file_path: Path to CV file
            
        Returns:
            Extracted text
        
        Raises:
            ValueError: If file format is not supported
        """
        if not os.path.exists(file_path):
            raise ValueError(f"File not found: {file_path}")
        
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        if ext == '.pdf':
            return CVParser.extract_text_from_pdf(file_path)
        elif ext == '.docx':
            return CVParser.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}. Supported formats: {CVParser.SUPPORTED_FORMATS}")
